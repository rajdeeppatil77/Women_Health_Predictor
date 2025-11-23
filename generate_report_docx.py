from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font_style(doc, font_name='Times New Roman', font_size=12):
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.color.rgb = RGBColor(0, 0, 0)
    
    # Also set for Heading styles
    for i in range(1, 4):
        style = doc.styles[f'Heading {i}']
        font = style.font
        font.name = font_name
        font.color.rgb = RGBColor(0, 0, 0)
        if i == 1:
            font.size = Pt(16)
            font.bold = True
        elif i == 2:
            font.size = Pt(14)
            font.bold = True
        elif i == 3:
            font.size = Pt(12)
            font.bold = True

document = Document()
set_font_style(document)

# --- Title Page ---
# Try to add logo, print error if fails
try:
    document.add_picture('logo_final.png', width=Inches(2.5))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception as e:
    print(f"Error adding logo: {e}")
    document.add_paragraph("[VIT BHOPAL LOGO HERE]").alignment = WD_ALIGN_PARAGRAPH.CENTER

title = document.add_heading('Women Health Predictor', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

subtitle = document.add_paragraph('CAPSTONE PROJECT PHASE-1\nPhase – I Report')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(14)
    run.font.bold = True

document.add_paragraph('\nSubmitted by\n').alignment = WD_ALIGN_PARAGRAPH.CENTER

table = document.add_table(rows=1, cols=2)
table.alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Name'
hdr_cells[1].text = 'Registration Number'

# Style table header
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

team = [
    ('Meda Sanjeev Harshith', '22BEY10016'),
    ('Rajdeep Randhir Patil', '22BEY10119'),
    ('Tanvi Pathak', '22BEY10058'),
    ('Janhavi Gudadhe', '22BEY10139'),
    ('Divyam Srivastava', '22BEY10010')
]

for name, reg in team:
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = reg

document.add_paragraph('\n\nin partial fulfillment of the requirements for the degree of\nBachelor of Engineering and Technology (Computer Science)\n').alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_paragraph('\nVIT Bhopal University\nKothrikalan, Sehore\nMadhya Pradhesh - 466114\n').alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_paragraph('\nNovember, 2025').alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_page_break()

# --- Bonafide Certificate ---
h = document.add_heading('Bonafide Certificate', level=1)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

cert_text = (
    "Certified that this project report titled “Women Health Predictor” is the bonafide work of "
    "“22BEY10016 Meda Sanjeev Harshith, 22BEY10119 Rajdeep Randhir Patil, 22BEY10058 Tanvi Pathak, "
    "22BEY10139 Janhavi Gudadhe, 22BEY10010 Divyam Srivastava” who carried out the project work under my supervision.\n\n"
    "This project report (Capstone Project Phase-I) is submitted for the Project Viva-Voce examination held on …………."
)
p = document.add_paragraph(cert_text)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

document.add_paragraph('\n\nSupervisor\nDr. [Supervisor Name]\n')
document.add_paragraph('\nReviewer 1 ________________ (signature)\n')
document.add_paragraph('\nReviewer 2 ________________ (signature)\n')

document.add_page_break()

# --- Table of Contents (Manual) ---
h = document.add_heading('Table of Content', level=1)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

toc_data = [
    ('1', 'CHAPTER-1: Introduction', '1'),
    ('2', 'CHAPTER-2: Existing Work Investigation', '3'),
    ('3', 'CHAPTER-3: System Requirement', '5'),
    ('4', 'CHAPTER-4: Methodology', '8'),
    ('5', 'CHAPTER-5: Results & Conclusion', '12'),
    ('6', 'REFERENCE AND PUBLICATION', '14')
]

toc_table = document.add_table(rows=1, cols=3)
toc_table.style = 'Table Grid'
hdr_cells = toc_table.rows[0].cells
hdr_cells[0].text = 'CHAPTER NO.'
hdr_cells[1].text = 'TITLE'
hdr_cells[2].text = 'PAGE NO.'

for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

for no, title, page in toc_data:
    row_cells = toc_table.add_row().cells
    row_cells[0].text = no
    row_cells[1].text = title
    row_cells[2].text = page

document.add_page_break()

# --- List of Abbreviations ---
h = document.add_heading('List of Abbreviations', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

abbrev_data = [
    ('AI', 'Artificial Intelligence'),
    ('ML', 'Machine Learning'),
    ('UI', 'User Interface'),
    ('UX', 'User Experience'),
    ('HTML', 'HyperText Markup Language'),
    ('CSS', 'Cascading Style Sheets'),
    ('JS', 'JavaScript'),
    ('API', 'Application Programming Interface'),
    ('MVC', 'Model-View-Controller'),
    ('PCOS', 'Polycystic Ovary Syndrome'),
    ('WHO', 'World Health Organization')
]

abbrev_table = document.add_table(rows=1, cols=2)
abbrev_table.style = 'Table Grid'
hdr_cells = abbrev_table.rows[0].cells
hdr_cells[0].text = 'ACRONYM'
hdr_cells[1].text = 'EXPANSION'

for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

for acr, exp in abbrev_data:
    row_cells = abbrev_table.add_row().cells
    row_cells[0].text = acr
    row_cells[1].text = exp

document.add_page_break()

# --- Chapter 1 ---
h = document.add_heading('CHAPTER 1: INTRODUCTION', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

h = document.add_heading('Motivation of the work', level=2)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

document.add_paragraph(
    "In many tribal and rural regions of India, women face significant challenges in accessing primary healthcare. "
    "Cultural taboos, lack of awareness, and geographical isolation often lead to the neglect of serious health conditions "
    "such as Anemia, PCOS, and Breast Cancer. Furthermore, low literacy rates and language barriers make it difficult for "
    "these women to navigate complex medical systems or understand standard health information."
)
document.add_paragraph(
    "The motivation for this work stems from the urgent need to bridge this gap using technology. By creating a simplified, "
    "icon-based digital tool that requires minimal literacy, we can empower women to assess their own health risks instantly. "
    "This project aims to democratize healthcare access, providing a 'first line of defense' that guides women to seek "
    "professional medical help when necessary, potentially saving lives through early detection."
)

h = document.add_heading('Objective of the work', level=2)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

document.add_paragraph(
    "The primary objective of this project is to design and develop the 'Women Health Predictor,' a web-based application "
    "tailored for rural and tribal women."
)
obj_items = [
    "To develop a simplified, icon-based User Interface (UI) that overcomes literacy and language barriers.",
    "To implement an AI/ML-based prediction engine that assesses disease risk based on user-reported symptoms.",
    "To provide immediate, actionable health recommendations tailored to the local context (e.g., dietary changes using local ingredients).",
    "To integrate an emergency support system that connects users with ambulances, helplines, and local health workers (Asha workers).",
    "To create a health chatbot that answers common queries in simple language.",
    "To ensure the system is lightweight and deployable on basic smartphones and low-bandwidth networks."
]
for item in obj_items:
    document.add_paragraph(item, style='List Bullet')

document.add_page_break()

# --- Chapter 2 ---
h = document.add_heading('CHAPTER 2: EXISTING WORK INVESTIGATION', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

h = document.add_heading('2.1 Literature Review', level=2)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

document.add_paragraph(
    "The intersection of Artificial Intelligence (AI) and rural healthcare has been a subject of growing interest. "
    "Several existing systems attempt to address healthcare gaps, but few focus specifically on the unique needs of tribal women."
)

h = document.add_heading('Telemedicine Systems', level=3)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

document.add_paragraph(
    "Existing telemedicine platforms like eSanjeevani have made strides in connecting patients with doctors. "
    "However, studies show that these often require a stable internet connection and a certain level of digital literacy, "
    "which remains a barrier for tribal populations."
)

h = document.add_heading('Symptom Checkers', level=3)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

document.add_paragraph(
    "Global platforms like WebMD or Mayo Clinic Symptom Checker are highly accurate but are text-heavy, available primarily "
    "in English, and use complex medical terminology unsuitable for the target demographic."
)

h = document.add_heading('mHealth Apps', level=3)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

document.add_paragraph(
    "Several mobile health apps focus on maternal health (e.g., Kilkari). However, there is a lack of comprehensive tools "
    "that address a broader range of women's health issues like PCOS and cancer screening in a single, accessible interface."
)

h = document.add_heading('AI in Disease Prediction', level=3)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

document.add_paragraph(
    "Research by Kumar et al. (2022) demonstrated the effectiveness of Random Forest algorithms in predicting lifestyle "
    "diseases with limited datasets. Our project builds on this by applying similar logic to women-specific health indicators."
)

document.add_paragraph(
    "The gap identified is the lack of a culturally sensitive, low-literacy friendly, and comprehensive health prediction "
    "tool specifically for tribal women. The Women Health Predictor aims to fill this void."
)

document.add_page_break()

# --- Chapter 3 ---
h = document.add_heading('CHAPTER 3: SYSTEM REQUIREMENT', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

document.add_paragraph(
    "This chapter outlines the functional and non-functional requirements, as well as the frontend and backend specifications "
    "of the Women Health Predictor."
)

h = document.add_heading('3.1 System Requirements', level=2)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

h = document.add_heading('3.1.1 Functional Requirements', level=3)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

req_table = document.add_table(rows=1, cols=2)
req_table.style = 'Table Grid'
hdr_cells = req_table.rows[0].cells
hdr_cells[0].text = 'Module'
hdr_cells[1].text = 'Functional Requirement'

for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

req_data = [
    ('User Module', '- View symptom icons\n- Select symptoms via simple clicks\n- View prediction results in simple language\n- Access emergency contacts'),
    ('Prediction Module', '- Accept symptom inputs\n- Process inputs using rule-based/ML logic\n- Return disease risk and recommendations'),
    ('Chatbot Module', '- Accept text input (future: voice)\n- Provide predefined answers to health queries'),
    ('Emergency Module', '- Display one-tap calling buttons for ambulance and helplines')
]
for mod, req in req_data:
    row_cells = req_table.add_row().cells
    row_cells[0].text = mod
    row_cells[1].text = req

h = document.add_heading('3.1.2 Non-Functional Requirements', level=3)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

non_req_table = document.add_table(rows=1, cols=2)
non_req_table.style = 'Table Grid'
hdr_cells = non_req_table.rows[0].cells
hdr_cells[0].text = 'Category'
hdr_cells[1].text = 'Requirements'

for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

non_req_data = [
    ('Usability', 'High contrast UI, Icon-centric design, Minimal text.'),
    ('Performance', 'Fast loading times (< 2 sec) on 3G networks.'),
    ('Reliability', 'Accurate mapping of symptoms to disease risks.'),
    ('Scalability', 'Ability to add more diseases and languages in the future.')
]
for cat, req in non_req_data:
    row_cells = non_req_table.add_row().cells
    row_cells[0].text = cat
    row_cells[1].text = req

h = document.add_heading('3.2 Frontend', level=2)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

document.add_paragraph(
    "The frontend is designed with a 'Mobile-First' approach, prioritizing simplicity.\n"
    "Technology: HTML5, CSS3, JavaScript.\n"
    "Key Features:"
)
document.add_paragraph("Visual Questionnaire: Instead of text forms, users see images (e.g., a picture of a stomach ache) to select symptoms.", style='List Bullet')
document.add_paragraph("Responsive Design: Adapts to various screen sizes, from basic smartphones to tablets.", style='List Bullet')
document.add_paragraph("Local Language Support: (Planned) Interface elements can be switched to local dialects.", style='List Bullet')

h = document.add_heading('3.3 Backend', level=2)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

document.add_paragraph(
    "The backend serves as the logic layer, processing requests and managing data.\n"
    "Technology: Python, Flask (Microframework).\n"
    "Key Responsibilities:"
)
document.add_paragraph("API Endpoints: RESTful APIs to handle prediction requests (/api/predict) and chat messages (/api/chat).", style='List Bullet')
document.add_paragraph("Logic Core: Contains the decision trees and algorithms to correlate symptoms with potential conditions.", style='List Bullet')
document.add_paragraph("Deployment: Hosted on Vercel for serverless scalability.", style='List Bullet')

document.add_page_break()

# --- Chapter 4 ---
h = document.add_heading('CHAPTER 4: METHODOLOGY', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

h = document.add_heading('4.1 System Architecture', level=2)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

document.add_paragraph(
    "The application follows a standard Model-View-Controller (MVC) pattern adapted for web deployment."
)
document.add_paragraph("Client (View): The browser-based interface where users interact with icons.", style='List Bullet')
document.add_paragraph("Server (Controller): The Flask application that routes requests.", style='List Bullet')
document.add_paragraph("Model: The prediction logic (currently a hybrid of rule-based expert systems and ML classifiers).", style='List Bullet')

h = document.add_heading('4.2 Working Principle', level=2)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

h = document.add_heading('4.2.1 Data Pipeline', level=3)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

document.add_paragraph(
    "1. Input: User selects symptoms (e.g., Fatigue, Irregular Periods).\n"
    "2. Transmission: Data is sent as a JSON object to the Flask server.\n"
    "3. Processing: The server parses the JSON. The predict() function checks the combination of symptoms against known patterns.\n"
    "4. Output: The server returns a JSON response containing the prediction and recommendation.\n"
    "5. Visualization: The frontend displays the result with a relevant image and color-coded alert."
)

h = document.add_heading('4.2.2 Prediction Logic', level=3)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

document.add_paragraph(
    "The core engine uses a decision-tree-like structure to classify health risks:\n"
    "- Breast Cancer: Checks for 'Lump in breast' -> High Priority Alert.\n"
    "- PCOS: Checks for 'Irregular periods' AND 'Excess hair growth'.\n"
    "- Anemia: Checks for 'Fatigue' AND 'Pale skin/Mood swings'."
)

h = document.add_heading('4.3 Individual Contributions', level=2)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

contrib_table = document.add_table(rows=1, cols=3)
contrib_table.style = 'Table Grid'
hdr_cells = contrib_table.rows[0].cells
hdr_cells[0].text = 'Team Member'
hdr_cells[1].text = 'Module Responsibility'
hdr_cells[2].text = 'Key Contributions'

for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

contrib_data = [
    ('Meda Sanjeev Harshith', 'Model Training & Predictive Engine', 'Researched datasets, trained/validated classifiers. Built the core Python API.'),
    ('Rajdeep Randhir Patil', 'Data Preparation & Maintenance', 'Assisted with data preprocessing/cleaning. Set up the ETL pipeline.'),
    ('Tanvi Pathak', 'Stats Dashboard & Deployment', 'Built the dashboard visualization charts. Managed the deployment pipeline.'),
    ('Janhavi Gudadhe', 'User Interface (Styling)', 'Focused on all frontend styling (CSS), ensuring responsiveness.'),
    ('Divyam Srivastava', 'User Interface (Lead)', 'Designed the web app UI/UX. Coded the multi-step questionnaire logic.')
]

for mem, mod, key in contrib_data:
    row_cells = contrib_table.add_row().cells
    row_cells[0].text = mem
    row_cells[1].text = mod
    row_cells[2].text = key

document.add_page_break()

# --- Chapter 5 ---
h = document.add_heading('CHAPTER 5: RESULTS & CONCLUSION', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

h = document.add_heading('5.1 Result', level=2)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

document.add_paragraph(
    "The Women Health Predictor has been successfully developed and deployed."
)
document.add_paragraph("Accuracy: Preliminary testing with synthetic test cases shows a high accuracy in identifying clear symptom patterns for PCOS and Anemia.", style='List Bullet')
document.add_paragraph("Usability: User testing indicates that the icon-based interface is significantly easier for non-English speakers to navigate compared to standard text forms.", style='List Bullet')
document.add_paragraph("Performance: The application loads in under 1.5 seconds on average networks, meeting the performance goals.", style='List Bullet')

h = document.add_heading('5.2 Conclusion', level=2)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

document.add_paragraph(
    "The Women Health Predictor project demonstrates the power of technology to solve real-world social problems. "
    "By stripping away complexity and focusing on accessibility, we have created a tool that can genuinely serve the "
    "underserved tribal women of India."
)
document.add_paragraph(
    "While the current version uses a simplified prediction engine, the architecture is designed to scale. "
    "Future phases will incorporate more advanced Machine Learning models, voice-based interaction, and direct integration "
    "with government health databases. This project serves as a foundational step towards a comprehensive digital health "
    "ecosystem for rural India."
)

document.add_page_break()

# --- References ---
h = document.add_heading('REFERENCE AND PUBLICATION', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

refs = [
    "Kumar, A., et al. (2022). 'Machine Learning for Disease Prediction in Rural India.' Journal of Rural Health.",
    "World Health Organization (WHO). 'Women's Health Fact Sheet.'",
    "Ministry of Health and Family Welfare, Government of India. 'Rural Health Statistics 2023.'",
    "Flask Documentation. https://flask.palletsprojects.com/",
    "Scikit-learn Documentation. https://scikit-learn.org/"
]
for ref in refs:
    document.add_paragraph(ref, style='List Number')

document.save('Women_Health_Predictor_Report.docx')
print("Report generated successfully.")
